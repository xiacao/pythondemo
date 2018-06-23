from docx import Document
from docx.shared import Inches

document = Document()#创建文档
document.add_heading('Document Title', 0)#大标题等级为0

p = document.add_paragraph('A plain paragraph having some')#添加文本段落
p.add_run('bold').bold = True#添加文本为粗体
p.add_run(' and some ')
p.add_run('italic.').italic = True#添加文本为斜体

document.add_heading('Heading, level 1', level = 1)
document.add_paragraph('Intense quote', style='IntenseQuote')#添加文本段落样式为Style
document.add_paragraph('first item in unordered list', style='ListBullet')
document.add_paragraph('fisrt item in ordered list', style='ListNumber')
document.add_picture('python_logo.gif', width=Inches(1.25))#添加图片宽度为1.25英寸

table = document.add_table(rows = 1, cols = 3)#创建了一行三列的表格
hdr_cells = table.rows[0].cells#取得这个表格的第一行
hdr_cells[0].text = 'Qty'#为列名赋值
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for item in range(3):
    row_cells = table.add_row().cells#添加新行
    row_cells[0].text = str(item)#为每列赋值
    row_cells[1].text = str(item*100)
    row_cells[2].text = str(item)+'10'

document.add_page_break()#添加分页符到文档末尾
document.save('hello.docx')#保存文档
